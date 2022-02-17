# -*- coding: utf-8 -*-

import sys
from datetime import datetime
import os
import pathlib as path
import pandas as pd
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from bin.load_data import load_and_preprocessing_data, read_config_and_language

#####################################################
#                                                   #
#                General logic                      #
#                                                   #
#####################################################
CONFIG, LANGUAGE = read_config_and_language()

# ToDo Prescribe the system requirements for the program
#  (Windows 7 is not suitable)


def combine_word_documents(based_doc, path_merged_doc):
    """Copies the entire document from path_marged_doc to the end of based_doc
    Formatting of the copied elements is preserved if possible

    :param based_doc: The main doc at the end of which the new info is copied
    :type based_doc: docx.document.Document
    :param path_merged_doc: The path (or object) to the doc for copy
    :type path_merged_doc: str | docx.document.Document
    :return: Glued Word document
    :rtype: docx.document.Document

    """
    if type(path_merged_doc) is str:
        merged_doc = docx.Document(path_merged_doc)
    else:
        merged_doc = path_merged_doc

    for element in merged_doc.element.body:
        based_doc.element.body.append(element)

    return based_doc


def save_to_docx(document, path_dir_to_save, filename_key):
    """Saves the Word document (AA or UI) to the desired directory with an
    indication of the creation date

    :param document: The document to be saved
    :type document: docx.document.Document
    :param path_dir_to_save: The path to the directory to save the word doc to
    :type path_dir_to_save: str
    :param filename_key: A key file name from configs
    :type filename_key: str
    :return: Just saves
    :rtype: None

    """
    # Save the newly created file in .docx format
    created_date = datetime.now().strftime('(%Y-%m-%d_%H-%M)')
    docx_name = CONFIG.get(LANGUAGE, filename_key)
    document.save(f"{path_dir_to_save}{os.sep}{docx_name}{created_date}.docx")


#####################################################
#                                                   #
#                   UI LOGIC                        #
#                                                   #
#####################################################
def create_df_ui(df_authors):
    """Creates a dataframe for the table in the Notification of Performers:
    | № | Full name | Place of work and position | Consent |

    :param df_authors: Source dataframe with all data on authors
    :type df_authors: pandas.core.frame.DataFrame
    :return: Dataframe for notifying performers
    :rtype: pandas.core.frame.DataFrame

    """
    try:
        # Initial an empty dataframe create uvedomlenie ispolniteley
        finish_df = pd.DataFrame()

        # Create the first column "serial number"
        finish_df['number'] = pd.Series(range(1, len(df_authors) + 1)).astype(
            str)

        # Create the full names of authors (ФИО)
        finish_df['full_name'] = df_authors['last_name'] + "\n" + \
                                 df_authors['first_name'] + "\n" + \
                                 df_authors['middle_name']

        # work_place is [organisation,\n post,\n academic]
        finish_df['work_place'] = df_authors['job'] + ",\n" + \
                                  df_authors['post'] + ",\n" + \
                                  df_authors['academic']
        # Delete marked absences of academic rank
        finish_df['work_place'] = finish_df['work_place'].apply(
            delete_final_comma)

        return finish_df
    except KeyError:
        raise KeyError(sys.exc_info())


def delete_final_comma(element):
    """Remove the last "empty" comma

    :param element: Element from the 'work_space' column
    :type element: str
    :return: Correct value, without extra characters
    :rtype: str

    """
    # Remove leading and ending whitespace characters from name of the columns
    element = str(element).strip()

    # Find and remove the last and useless comma from the end of the line
    if element[-1] == ',':
        element = element[:-1]

    return element


def fill_ui_table(df_ui):
    """Open the 1-th part of UI document (.docx) and fill the table in this doc

    :param df_ui: Dataframe with special processing info about authors for UI
    :type df_ui: pandas.core.frame.DataFrame
    :return: The 1-th part of UI for saving
    :rtype: docx.document.Document

    """
    # Open of template for ui
    document = docx.Document(f'{path.Path.cwd()}{os.sep}resources{os.sep}'
                             f'templates{os.sep}ui_part_1.docx')

    # Take the last table (1.5 Authors table with column names and 1 empty row)
    table = document.tables[-1]

    # Saving the number of rows in the table
    n_authors = df_ui.shape[0]

    # Add the missing rows in the table
    for _ in range(n_authors - 1):
        table.add_row()

    # Create the approval column by default
    approval = 'Не требуется'

    # Fill the table
    for n_author in range(n_authors):
        for col in range(4):
            # Get a cell from the table
            cell = table.cell(n_author + 1, col)

            # Write our data into cell
            if col == 3:
                cell.text = approval
            else:
                cell.text = df_ui.iloc[n_author, col]

            # Center alignment of text inside a special cell (first and last)
            if col in (0, 3):
                cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    return document


def generate_file_ui(df_authors, dir_for_save=''):
    """The finished function for the full cycle of UI generation

    :param df_authors: A row dataframe with info about authors
    :type df_authors: pandas.core.frame.DataFrame
    :param dir_for_save: The path to the directory to save the created UI
    :type dir_for_save: str
    :return: Save table for UI document (.docx)
    :rtype: None

    """
    # Processing the df_input for table in UI.docx document
    df_UI = create_df_ui(df_authors)

    # Open template and fill the table
    document = fill_ui_table(df_UI)

    # Finishing of document
    path_merged_doc = f'{path.Path.cwd()}{os.sep}resources{os.sep}templates{os.sep}ui_finish.docx'
    document = combine_word_documents(document, path_merged_doc)

    # Save the document
    save_to_docx(document=document,
                 path_dir_to_save=dir_for_save,
                 filename_key='-filename_UI-')


#####################################################
#                                                   #
#                   UA LOGIC                        #
#                    part1                          #
#####################################################
def create_df_ua_part1(df_authors):
    """Preparing a dataframe for creating the first part of the UA
    (long text for each author)

    :param df_authors: The original dataframe with all data about the authors
    :type df_authors: pandas.core.frame.DataFrame
    :return: Dataframe with info specifically for the first part of the UA
    :rtype: pandas.core.frame.DataFrame

    """
    try:
        finish_df = pd.DataFrame()

        # Full name for the first string
        finish_df['full_name'] = df_authors['last_name'] + ' ' + \
                                 df_authors['first_name'] + ' ' + \
                                 df_authors['middle_name']

        # Name and number of contract
        finish_df['contract'] = df_authors['contract']

        # Contribution of each authors in the total result
        finish_df['contribution'] = df_authors['contribution']

        # It's a datetime column (employment contract or not)
        finish_df['date_employ'] = pd.to_datetime(df_authors['date_employ'])

        # TRINITI employee or not (True if works in TRINITI)
        # ToDo Fill in the fields with True/False values (pay attention to the
        #  comparison)
        finish_df['TRINITI_employee'] = [finish_df['date_employ'] is not pd.NaN]

        return finish_df
    except KeyError:
        raise KeyError(sys.exc_info())


def make_table_for_one(df_row, based_doc):
    """Adding to the document (UA.docx ) the text part for the author from this
    df_row

    :param df_row: A string from a dataframe with information from one author
    :type df_row: pandas.core.series.Series
    :param based_doc: The document in which to write the text part by the author
    :type based_doc: docx.document.Document
    :param path_table_for_one: The path to the doc with the table of template
                               for one author
    :type path_table_for_one: str
    :param organization: The organization that we are notifying
    :type organization: str
    :return: The document with added text information about the author
    :rtype: docx.document.Document

    """
    print(f"\tmake_table_for_one\t\tdf_row['date_employ'] = {df_row['date_employ']}")
    if df_row['TRINITI_employee']:
        path_table_for_one = f'{path.Path.cwd()}{os.sep}resources' \
                             f'{os.sep}templates{os.sep}ua_table_for_one_from_TRINITI.docx'
    else:
        path_table_for_one = f'{path.Path.cwd()}{os.sep}resources' \
                             f'{os.sep}templates{os.sep}ua_table_for_one_not_from_TRINITI.docx'

    # Adding a table to fill in the author's data
    based_doc = combine_word_documents(based_doc, path_table_for_one)

    # We get a table in which to fill in all the information
    tabel = based_doc.tables[-1]

    # Filling in the name
    tabel.cell(0, 1).text = df_row['full_name']

    # ToDo Check this formatting (employ date)
    # Fill the employ date for authors from TRINITI
    if df_row['TRINITI_employee']:

        paragraph = tabel.cell(5, 2).paragraphs[0]
        employ_date = df_row['date_employ']

        print(f'\t\tmake_table_for_one\t\t"%d" = {str(employ_date.strftime("%d"))}'
              f'\t"%m" = {str(employ_date.strftime("%m"))}'
              f'\t"%y" = {str(employ_date.strftime("%y"))}')

        # Write a day
        run_5_1 = paragraph.add_run()
        run_5_1.text = str(employ_date.strftime("%d"))
        run_5_1.font.underline = True

        # Write the space between the day and month
        paragraph.add_run().text = "» "

        # Write a month
        run_5_3 = paragraph.add_run()
        run_5_3.text = str(employ_date.strftime("%m"))
        run_5_3.font.underline = True

        # Write the space between the month and year
        paragraph.add_run().text = " 20"

        # Write the year
        run_5_5 = paragraph.add_run()
        run_5_5.text = str(employ_date.strftime("%y"))
        run_5_5.font.underline = True

        paragraph.add_run().text = " г."

    # Fill contract
    tabel.cell(11, 0).text = df_row['contract']

    # Fill contribution
    tabel.cell(17, 1).text = df_row['contribution']

    # Single indent after the table
    based_doc.add_paragraph()

    return based_doc


def fill_part_1_of_ua(df_authors):
    """Creating the text part of the UA in a doc file, by all authors

    :param df_authors: Prepared dataframe with authors data for the text block
    :type df_authors: pandas.core.frame.DataFrame
    :return: Doc file with the finished first (text) part of the UA
    :rtype: docx.document.Document

    """
    df = create_df_ua_part1(df_authors)

    ua_part_1 = docx.Document(f'{path.Path.cwd()}{os.sep}resources'
                              f'{os.sep}templates{os.sep}ua_part_1.docx')

    for index in df.index:
        ua_part_1 = make_table_for_one(df_row=df.iloc[index],
                                       based_doc=ua_part_1,)

    return ua_part_1


#####################################################
#                                                   #
#                   UA LOGIC                        #
#                    part2                          #
#####################################################
def create_df_ua_part2(df_authors):
    """Preparing a dataframe for creating the second part (table) of the UA

    :param df_authors: The original dataframe with all data about the authors
    :type df_authors: pandas.core.frame.DataFrame
    :return: Dataframe with info specifically for the second part of the UA
    :rtype: pandas.core.frame.DataFrame

    """
    # Initial an empty dataframe create UI
    try:
        finish_df = pd.DataFrame()
        finish_df['name'] = df_authors['last_name'] + ' ' + \
                            [x[:1] for x in df_authors['first_name']] + '.' + \
                            [x[:1] for x in df_authors['middle_name']] + '.'
        # Remove double dot for events when not exist a middle name
        finish_df['name'] = finish_df['name']. \
            apply(lambda x: str(x).replace('..', '.'))

        finish_df['date_UA'] = pd.to_datetime(df_authors['date_UA'])

        return finish_df

    except KeyError:
        raise KeyError(sys.exc_info())


def fill_part_2_of_ua(df_ua):
    """Add the table in UA and save the result UA document (both text and table)

    :param df_ua: Prepared dataframe with authors data for the table
    :type df_ua: pandas.core.frame.DataFrame
    :return: The MS doc with the second part of ua
    :rtype: docx.document.Document

    """
    # Open the second part of ua
    doc = docx.Document(f'{path.Path.cwd()}{os.sep}resources{os.sep}'
                        f'templates{os.sep}ua_part_2.docx')

    # Calculate a shape of the table
    count_authors = df_ua.shape[0]
    table_rows = count_authors * 2
    table_cols = 3

    # Create table
    table = doc.add_table(table_rows, table_cols)

    # Fill the table with the data
    for row in range(table_rows):
        # Short cell names of one row of the table
        cell_0 = table.cell(row, 0)
        cell_1 = table.cell(row, 1)
        cell_2 = table.cell(row, 2)

        if row % 2 == 0:
            # ToDo Check underscore width - normal or not
            # Cell for signature
            cell_0.paragraphs[0].add_run().text = '______________/'

            # Cell with author name
            run_1 = cell_1.paragraphs[0].add_run()
            run_1.text = df_ua['name'].iloc[int(row/2)]
            run_1.font.underline = True

            # Cell with date of signature
            run_2 = cell_2.paragraphs[0].add_run()
            run_2.text = create_sign_date(date=df_ua['date_UA'].iloc[int(row/2)])

        else:
            # Explanations under the data
            cell_0.paragraphs[0].add_run().text = '(подпись)'
            cell_1.paragraphs[0].add_run().text = '(Фамилия И. О.)'

    # Applying styles to a table (For even and odd rows, then for the header)
    # ToDo Options for applying styles to the table:
    #  A. Apply styles to even and odd lines (will the library allow?)
    #  B. Apply a general style (for the main text (even lines) and manually:
    #  set the size to odd lines, underline the last name

    return doc



def create_sign_date(date):
    """Creating a line with the signature date in the desired design

    :param date: The date of signing of the UA, If there is no value,
                 then a mask is created to enter the date manually
    :type date: pandas._libs.tslibs.timestamps.Timestamp
    :return: The date in the desired format for insertion into the table
    :rtype: str

    """
    if date is pd.NaT:
        sign_date = f'Дата: «__» ____________ 20__г.'

    else:
        # ToDo A month to write letters in Russian (Maybe '%B')
        # ToDo Check the dependence on the global PC locale
        #  (Is it always in Russian?)
        sign_date = f'Дата: «{date.strftime("%d")}»' \
                    f' {date.strftime("%B")} {date.year}г.'

    return sign_date


def generate_file_ua(df_authors, dir_for_save=''):
    """Creating the UA doc file based on the authors' data

    :param df_authors: The original dataframe with all the info about the authors
    :type df_authors: pandas.core.frame.DataFrame
    :param dir_for_save: The path to the directory to save the created file
    :type dir_for_save: str
    :return: Save the created file
    :rtype: None

    """
    # Create doc with text part of UA
    part_1_of_ua = fill_part_1_of_ua(df_authors)

    # Adding the table create_df_ua_part2(df_in)
    part_2_of_ua = fill_part_2_of_ua(create_df_ua_part2(df_authors))

    finish_doc = combine_word_documents(part_1_of_ua, part_2_of_ua)
    finish_doc = combine_word_documents(finish_doc,
                                        f'{path.Path.cwd()}{os.sep}resources'
                                        f'{os.sep}templates{os.sep}ui_finish.docx')

    save_to_docx(finish_doc, dir_for_save)


#####################################################
#                                                   #
#                  MAKE LOGIC                       #
#                                                   #
#####################################################
def make_ui_and_ua(path_to_authors_data, dir_for_save_file):
    """Launch creating UI and UA

    :param path_to_authors_data: The path to the source data about the authors
    :type path_to_authors_data: str
    :param dir_for_save_file: The path to the directory to save the created docs
    :type dir_for_save_file: str
    :return: Saved doc files
    :rtype: None

    """
    df_in = load_and_preprocessing_data(path_to_authors_data)

    generate_file_ui(df_in, dir_for_save=dir_for_save_file)
    generate_file_ua(df_in, dir_for_save=dir_for_save_file)


def make_only_ui(path_to_authors_data, dir_for_save_file):
    """Launch creating only UI

    :param path_to_authors_data: The path to the source data about the authors
    :type path_to_authors_data: str
    :param dir_for_save_file: The path to the directory to save the created file
    :type dir_for_save_file: str
    :return: Saved a doc file
    :rtype: None

    """
    df_in = load_and_preprocessing_data(path_to_authors_data)

    generate_file_ui(df_in, dir_for_save_file)


def make_only_UA(path_to_authors_data, dir_for_save_file):
    """Launch creating only UA

    :param path_to_authors_data: The path to the source data about the authors
    :type path_to_authors_data: str
    :param dir_for_save_file: The path to the directory to save the created file
    :type dir_for_save_file: str
    :return: Saved a doc file
    :rtype: None

    """
    df_in = load_and_preprocessing_data(path_to_authors_data)

    generate_file_ua(df_in, dir_for_save=dir_for_save_file)
